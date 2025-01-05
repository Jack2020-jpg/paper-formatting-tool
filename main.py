from gui import GUI
from tkinter import messagebox
import re
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import os
from docx import Document
from PIL import Image
from io import BytesIO
from time import sleep
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 字号和 磅值的对应关系
font_sizes = {
    '八号': Pt(5),
    '七号': Pt(5.5),
    '小六': Pt(6.5),
    '六号': Pt(7.5),
    '小五': Pt(9),
    '五号': Pt(10.5),
    '小四': Pt(12),
    '四号': Pt(14),
    '小三': Pt(15),
    '三号': Pt(16),
    '小二': Pt(18),
    '二号': Pt(22),
    '小一': Pt(24),
    '一号': Pt(26),
    '小初': Pt(36),
    '初号': Pt(42),
}


# 分析文档架构
def analyze_structure(para):
    """
    分析每段内容，分析其属于哪种类型
    :param para: 原文档句柄
    :return:
        0 - 非标题非说明性文字，原样写入？
        1 - 一级标题
        2 - 二级标题
        3 - 三级标题
        4 - 表格说明文字
        5 - 图片说明文字
        6 - 参考文献
    """
    # 匹配一级标题
    if re.match(r'^(第[一二三四五六七八九十]+章[：:]*|^\d+\s+).+',
                para) or para.replace(" ","") == "摘要" or para == "Abstract" or para.replace(" ","") == "致谢" or para == "参考文献" or para.strip() == "附录" or para.replace(" ","") == "附录软件源码":
        return 1

    # 匹配三级标题 (必须在二级标题之前匹配)
    elif re.match(r'^\d+\.\d+\.\d+(\.\d+)*\s*.+', para):
        return 3

    # 匹配二级标题
    elif re.match(r'^\d+\.\d+\s*.+', para):
        return 2

    # 匹配表格说明
    elif re.match(r'^表\d+\s+.+', para):
        return 4

    # 匹配图片说明
    elif re.match(r'^图\d+\s+.+', para):
        return 5

    # 匹配参考文献
    elif re.match(r'^\[\d+\].+', para):
        return 6

    return 0


# 将对应的样式写入到新文档中
def apply_style_to_paragraph(new_doc, text, word_size="小四", word_type="宋体", is_bold=False, align_style="left",
                             suojin_need=True, line_distance=20):
    """
    将指定样式的内容写入到word中
    :param new_doc:新文档的新段落的句柄
    :param text: 要写入的内容
    :param word_size: 字号
    :param word_type: 中文字体
    :param is_bold: 是否加粗
    :param align_style: 对其方式
    :param suojin_need: 是否需要缩进，默认情况下只需要一级标题不需要缩进
    :param line_distance:
    :return:
    """
    # 在新文档中创建段落
    new_para = new_doc.add_paragraph()  # 创建空白段落

    # 将文本添加到段落中
    run = new_para.add_run(text)
    paragraph_format = new_para.paragraph_format
    # 设置行间距,默认为 20磅
    paragraph_format.line_spacing = Pt(line_distance)
    # 设置段前段后间距为 0 磅
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    if suojin_need:
        paragraph_format.first_line_indent = 0
        # 设置缩进 为两个字符 - 200 ，一个字符就是 100
        paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '200')

    font = run.font
    # 西文 字体
    font.name = 'Times New Roman'
    # 中文字体
    font.element.rPr.rFonts.set(qn("w:eastAsia"), word_type)
    font.size = font_sizes[word_size]
    if is_bold:
        run.bold = True
    if align_style == "left":
        # 左对齐
        new_para.alignment = 0
    elif align_style == "center":
        # 居中对齐
        new_para.alignment = 1
    elif align_style == "right":
        # 居中对齐
        new_para.alignment = 2


def main():
    # 显示界面，获取用户选择的文件和设置
    gui = GUI()
    file_path = gui.run()

    # 检查文件是否选择
    if not file_path:
        messagebox.showerror("错误", "没有选择文件！")
        return

    # 获取用户输入的设置（字体和行间距）
    settings = gui.get_user_settings()

    # 读取文档内容
    doc = Document(file_path)

    # 创建一个新的文档来保存修改后的内容
    new_doc = Document()

    # 初始化图片计数器
    image_counter = 1

    # 定义命名空间
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    }

    for paragraph_index, para in enumerate(doc.paragraphs):
        # 清除原有的缩进
        para.paragraph_format.left_indent = None
        para.paragraph_format.right_indent = None
        para.paragraph_format.first_line_indent = None

        para_text = para.text.strip()
        # 空行，跳过
        if para.text.strip() != "":
            paragraph_type = analyze_structure(para_text)
            print(paragraph_type)
            # 默认情况 宋体 小四  不加粗 需要缩进 左对齐
            text_font, text_size, is_bold, suojin_need, align_style = "宋体", "小四", False, True, "left"

            # 这里要保证 图表说明部分 需要被处理
            if paragraph_type == 0 or (len(para_text) > 20 and "图" not in para_text) or (
                    len(para_text) > 20 and "表" not in para_text):
                pass
            else:
                if paragraph_type == 1:
                    try:
                        # 检查上一个段落是否是分页符
                        if new_doc.paragraphs[-1].text != "":
                            # 一级标题之前插入一个分页符
                            new_doc.add_page_break()
                    except:
                        print("当前内容为第一页，无需检测分页符，直接跳过即可！")
                    text_font = settings["heading1_font"]
                    text_size = settings["heading1_size"]
                    if settings["heading1_bold"] == "是":
                        is_bold = True
                    else:
                        is_bold = False
                    # 标题不要缩进
                    suojin_need = False
                    # 居中显示
                    align_style = "center"
                elif paragraph_type == 2:
                    text_font = settings["heading2_font"]
                    text_size = settings["heading2_size"]
                    if settings["heading2_bold"] == "是":
                        is_bold = True
                    else:
                        is_bold = False
                elif paragraph_type == 3:
                    text_font = settings["heading3_font"]
                    text_size = settings["heading3_size"]
                    if settings["heading3_bold"] == "是":
                        is_bold = True
                    else:
                        is_bold = False
                elif paragraph_type == 4:
                    text_font = settings["table_font"]
                    text_size = settings["table_size"]
                    if settings["table_bold"] == "是":
                        is_bold = True
                    else:
                        is_bold = False
                    # 居中显示
                    align_style = "center"
                elif paragraph_type == 5:
                    text_font = settings["image_font"]
                    text_size = settings["image_size"]
                    # 居中显示
                    align_style = "center"
                    if settings["image_bold"] == "是":
                        is_bold = True
                    else:
                        is_bold = False
                elif paragraph_type == 6:
                    text_font = settings["reference_font"]
                    text_size = settings["reference_size"]

            print(para_text, text_font, text_size, is_bold, align_style, suojin_need)
            apply_style_to_paragraph(new_doc, para_text, word_size=text_size, word_type=text_font, is_bold=is_bold,
                                     align_style=align_style, suojin_need=suojin_need)


        # 处理段落中的图片
        for run in para.runs:
            # 检查运行中是否包含图片
            drawing_elements = run.element.findall(".//w:drawing", namespaces=nsmap)
            if not drawing_elements:
                continue

            for drawing in drawing_elements:
                # 提取嵌入的图片关系 ID
                blip_elements = drawing.findall(".//a:blip", namespaces=nsmap)
                for blip in blip_elements:
                    embed_id = blip.attrib.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed_id and embed_id in doc.part.rels:
                        image_part = doc.part.rels[embed_id].target_part
                        image_data = image_part.blob
                        # 确定图片格式
                        image = Image.open(BytesIO(image_data))
                        image_format = image.format.lower()
                        # 图片保存的文件夹
                        output_folder = "tmp_img"
                        # 确保输出文件夹存在
                        os.makedirs(output_folder, exist_ok=True)

                        # 保存图片
                        image_filename = f"paragraph_{paragraph_index + 1}_image_{image_counter}.{image_format}"
                        image_path = os.path.join(output_folder, image_filename)
                        image.save(image_path)

                        # 等待图片保存结束
                        sleep(0.5)

                        # 将图片添加到新文件中
                        new_para = new_doc.add_paragraph()
                        new_para.alignment = 1

                        # 添加图片到新文档
                        run = new_para.add_run()
                        run.add_picture(image_path, width=Inches(4.25))

                        # Explicitly center the picture by setting the run alignment
                        run.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        print(f"已保存图片: {image_path}")
                        image_counter += 1



    output_file = 'modified_output1.docx'  # 保存修改后的文件路径
    messagebox.showinfo("完成", f"文件已保存为: {output_file}")
    new_doc.save(output_file)


if __name__ == "__main__":
    main()
